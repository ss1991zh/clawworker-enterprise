"""
明文文本附件提取 —— txt / md / docx / pdf / rtf / html。

附件文本会被折叠到 LLM 的 user prompt 顶部,不入加密管线。
"""

from __future__ import annotations

import re
from pathlib import Path
from typing import Optional


SUPPORTED_EXTS = {".txt", ".md", ".markdown", ".rst", ".log", ".text",
                  ".docx", ".pdf", ".rtf", ".html", ".htm", ".json", ".yml", ".yaml"}

# 安全上限:单文件提取后最多保留多少字符(防 prompt 爆)
MAX_TEXT_CHARS = 30_000


def is_text_attachment(filename: str) -> bool:
    ext = Path(filename).suffix.lower()
    return ext in SUPPORTED_EXTS


def extract(file_path: Path) -> str:
    """根据后缀分发提取器,返回纯文本(已 trim 到上限)。"""
    ext = file_path.suffix.lower()
    if ext in (".txt", ".md", ".markdown", ".rst", ".log", ".text",
               ".json", ".yml", ".yaml"):
        text = _read_plain(file_path)
    elif ext == ".docx":
        text = _read_docx(file_path)
    elif ext == ".pdf":
        text = _read_pdf(file_path)
    elif ext == ".rtf":
        text = _read_rtf(file_path)
    elif ext in (".html", ".htm"):
        text = _read_html(file_path)
    else:
        raise ValueError(f"不支持的文本附件格式:{ext}")

    text = text.strip()
    if len(text) > MAX_TEXT_CHARS:
        text = text[:MAX_TEXT_CHARS] + f"\n\n[文件过长,已截断 · 保留前 {MAX_TEXT_CHARS} 字]"
    return text


# -----------------------------------------------------------------------------
# 各格式提取器
# -----------------------------------------------------------------------------

def _read_plain(p: Path) -> str:
    """txt / md / log / json / yaml —— 多编码兜底。"""
    for enc in ("utf-8", "utf-8-sig", "gbk", "gb18030", "latin-1"):
        try:
            return p.read_text(encoding=enc)
        except UnicodeDecodeError:
            continue
    # 实在不行用 latin-1(永远不抛)
    return p.read_text(encoding="latin-1", errors="replace")


def _read_docx(p: Path) -> str:
    try:
        from docx import Document
    except ImportError as e:
        raise RuntimeError("缺少 python-docx,运行:pip install python-docx") from e
    doc = Document(str(p))
    parts: list[str] = []
    for para in doc.paragraphs:
        text = para.text
        if text:
            parts.append(text)
    # 表格也提一下
    for table in doc.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells]
            cells = [c for c in cells if c]
            if cells:
                parts.append(" | ".join(cells))
    return "\n".join(parts)


def _read_pdf(p: Path) -> str:
    try:
        from pypdf import PdfReader
    except ImportError as e:
        raise RuntimeError("缺少 pypdf,运行:pip install pypdf") from e
    reader = PdfReader(str(p))
    parts: list[str] = []
    for i, page in enumerate(reader.pages, 1):
        try:
            t = page.extract_text() or ""
        except Exception:
            t = ""
        if t.strip():
            parts.append(f"[Page {i}]\n{t.strip()}")
    return "\n\n".join(parts)


_RTF_CONTROL_RE = re.compile(r"\\[a-z]+-?\d*", re.IGNORECASE)
_RTF_BRACE_RE = re.compile(r"[{}]")


def _read_rtf(p: Path) -> str:
    """简易 RTF 去控制码 —— 满足 80% 的中英文文档。"""
    raw = _read_plain(p)
    raw = _RTF_CONTROL_RE.sub(" ", raw)
    raw = _RTF_BRACE_RE.sub("", raw)
    raw = re.sub(r"\s+", " ", raw).strip()
    return raw


_HTML_SCRIPT_RE = re.compile(r"<(script|style)[^>]*>.*?</\1>", re.IGNORECASE | re.DOTALL)
_HTML_TAG_RE = re.compile(r"<[^>]+>")
_HTML_ENTITY = {
    "&nbsp;": " ", "&amp;": "&", "&lt;": "<", "&gt;": ">",
    "&quot;": '"', "&#39;": "'",
}


def _read_html(p: Path) -> str:
    raw = _read_plain(p)
    raw = _HTML_SCRIPT_RE.sub("", raw)
    raw = _HTML_TAG_RE.sub(" ", raw)
    for k, v in _HTML_ENTITY.items():
        raw = raw.replace(k, v)
    raw = re.sub(r"[ \t]+", " ", raw)
    raw = re.sub(r"\n{3,}", "\n\n", raw)
    return raw.strip()
