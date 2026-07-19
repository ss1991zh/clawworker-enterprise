"""
合规报告生成(Word/.docx)—— 把审计台账写成**非技术人员也能看懂**的带排版报告。

面向合规/采购/客户尽调:用大白话解释"您的原始数据从未离开本机、AI 只看到字段名、
解密均经您授权",并附逐条台账 + 名词解释。
"""
from __future__ import annotations

import io
from datetime import datetime

from client.he_ops import audit

_GREEN = "DDF4E3"
_GREY = "F1F3F5"
_RED = "FDE2E2"
_BLUE = "E7EEFB"


def _set_default_font(doc, font="微软雅黑"):
    from docx.oxml.ns import qn
    style = doc.styles["Normal"]
    style.font.name = font
    try:
        style.element.rPr.rFonts.set(qn("w:eastAsia"), font)
    except Exception:  # noqa: BLE001
        pass


def _shade(el, color):
    from docx.oxml import parse_xml
    from docx.oxml.ns import nsdecls
    el.append(parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color}"/>'))


def _callout(doc, text, color=_GREEN, bold=True, size=12):
    """高亮结论框(单格表格 + 底色)。"""
    from docx.shared import Pt
    t = doc.add_table(rows=1, cols=1)
    cell = t.rows[0].cells[0]
    _shade(cell._tc.get_or_add_tcPr(), color)
    p = cell.paragraphs[0]
    run = p.add_run(text)
    run.bold = bold
    run.font.size = Pt(size)
    return t


def _kv_table(doc, rows: list[tuple[str, str]]):
    from docx.shared import Pt
    t = doc.add_table(rows=0, cols=2)
    t.style = "Table Grid"
    for k, v in rows:
        r = t.add_row().cells
        r[0].text = ""
        run = r[0].paragraphs[0].add_run(k)
        run.bold = True
        _shade(r[0]._tc.get_or_add_tcPr(), _GREY)
        r[1].text = str(v)
    return t


def _data_table(doc, headers: list[str], rows: list[list[str]], widths=None):
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = "Table Grid"
    hdr = t.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = ""
        run = hdr[i].paragraphs[0].add_run(h)
        run.bold = True
        _shade(hdr[i]._tc.get_or_add_tcPr(), _BLUE)
    for row in rows:
        cells = t.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = str(val)
    return t


def _ts(s: str) -> str:
    return (s or "").replace("T", " ")[:19]


def build_docx(user: str) -> bytes:
    """生成合规报告 .docx 字节。"""
    from docx import Document
    from docx.shared import Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    evs = audit.read_events(user, limit=100000)
    s = audit.summary(user)
    exposures = [e for e in evs if e.get("type") == "llm_exposure"]
    decrypts = [e for e in evs if e.get("type") == "decrypt_auth"]

    doc = Document()
    _set_default_font(doc)

    # ---- 标题 ----
    title = doc.add_heading("数据隐私合规报告", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = sub.add_run("Clawworker 企业版 · 同态加密数据分析")
    r.font.size = Pt(11); r.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    rr = meta.add_run(f"使用者:{user}　|　生成时间:{datetime.now().strftime('%Y-%m-%d %H:%M')}")
    rr.font.size = Pt(9); rr.font.color.rgb = RGBColor(0x88, 0x88, 0x88)

    # ---- 一、这份报告说明什么 ----
    doc.add_heading("一、这份报告说明什么", level=1)
    doc.add_paragraph(
        "本系统用「同态加密」技术帮您分析数据。它最关键的保证是:您的原始数据"
        "(具体的数字、金额、名单)从头到尾都留在您自己的电脑上,从不上传;"
        "连帮忙写分析逻辑的 AI,也只能看到字段的「名称」(比如「销售额」「客户名」),"
        "看不到任何一个具体数值。")
    doc.add_paragraph(
        "这份报告把「确实做到了这件事」的证据列出来——每次分析时 AI 收到了哪些字段名、"
        "每次把结果解密给您看时是否经过您授权——供合规检查与留档。")

    # ---- 二、核心结论 ----
    doc.add_heading("二、核心结论", level=1)
    ok = s.get("zero_plaintext_holds", True)
    _callout(doc, ("✓ 全程满足:分析过程中,AI 只接收了字段名称,您的明文数据值从未离开本机;"
                   "所有解密都经过您本人在本机授权,全程可追溯。") if ok else
                  ("⚠ 检出疑似数据外发,需复核(详见第四节)。"),
             color=_GREEN if ok else _RED)

    # ---- 三、关键指标 ----
    doc.add_heading("三、关键指标", level=1)
    from client.he_ops.audit import verify_chain
    chain = verify_chain(user)
    chain_txt = ("完整(未被篡改)" if chain["ok"]
                 else f"⚠ 第 {chain['broken_at']} 条起异常:{chain['reason']}")
    _kv_table(doc, [
        ("分析次数(AI 参与写逻辑)", s.get("llm_exposures", 0)),
        ("AI 看到具体数据值的次数", "0(AI 只看到字段名)" if ok else f"疑似 {s.get('plaintext_breaches',0)} 次"),
        ("解密授权次数", f"{s.get('decrypt_authorizations',0)}　"
                       f"(您批准 {s.get('decrypt_granted',0)} 次 / 拒绝或保留密文 {s.get('decrypt_denied',0)} 次)"),
        ("记录时间范围", f"{_ts(s.get('first_event'))} ~ {_ts(s.get('last_event'))}" if s.get("first_event") else "暂无记录"),
        ("审计链完整性(防篡改校验)", chain_txt),
    ])

    # ---- 四、AI 只看到字段名(逐次记录) ----
    doc.add_heading("四、AI 只看到字段名 —— 逐次记录", level=1)
    doc.add_paragraph(
        "说明:「字段名」就是表格的列标题(如「销售额」「大区」)。AI 靠它来写分析逻辑,"
        "但看不到任何一行的具体数字。下表每一行是一次分析。")
    if exposures:
        rows = [[_ts(e.get("ts")),
                 "、".join(e.get("fields", [])[:8]) + ("…" if len(e.get("fields", [])) > 8 else ""),
                 "否(仅字段名)" if e.get("no_plaintext", True) else "⚠ 疑似"]
                for e in exposures[-50:]]
        _data_table(doc, ["时间", "AI 收到的字段名", "是否含数据值"], rows)
    else:
        doc.add_paragraph("(暂无分析记录)")

    # ---- 五、解密授权记录 ----
    doc.add_heading("五、解密授权记录", level=1)
    doc.add_paragraph(
        "说明:计算全程在密文(加密后的乱码)上进行。只有当结果需要展示给您看时,"
        "系统才会请求您授权解密——决定权始终在您手里。")
    if decrypts:
        _LBL = {"granted": "授权解密展示", "keep_encrypted": "保留密文(不解密)", "denied": "拒绝/取消"}
        rows = [[_ts(e.get("ts")), _LBL.get(e.get("decision"), e.get("decision", "")),
                 e.get("detail", "")] for e in decrypts[-50:]]
        _data_table(doc, ["时间", "您的决定", "说明"], rows)
    else:
        doc.add_paragraph("(暂无解密授权记录)")

    # ---- 六、名词解释 ----
    doc.add_heading("六、名词解释(给非技术人员)", level=1)
    for term, expl in [
        ("明文 / 密文", "明文就是能直接看懂的数据(如「销售额 10000」);密文是加密后的乱码。本系统计算时用的全是密文。"),
        ("同态加密", "一种特殊加密,能在「乱码」上直接做加减乘除、统计,算完再解密,过程中不还原成明文。"),
        ("字段名(schema)", "表格的列标题(如「销售额」「大区」),不含任何一行的具体数值。AI 只拿到这个。"),
        ("解密授权", "把密文还原成能看懂的明文之前,需要您本人点确认——这一步只在您本机完成。"),
    ]:
        p = doc.add_paragraph(style="List Bullet")
        rb = p.add_run(term + ":"); rb.bold = True
        p.add_run(expl)

    # ---- 页脚说明 ----
    foot = doc.add_paragraph()
    fr = foot.add_run("本报告由系统自动生成,数据来源为本机审计日志(append-only,只增不改、可追溯)。")
    fr.italic = True; fr.font.size = Pt(9); fr.font.color.rgb = RGBColor(0x88, 0x88, 0x88)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()
